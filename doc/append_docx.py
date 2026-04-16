import docx
from docx import Document
import re
import os

template_path = r"d:\FPTU\SEM9\SEP490\FE\FDA-Web\doc\Template User Manual Report 6.docx"
out_path = r"d:\FPTU\SEM9\SEP490\FE\FDA-Web\doc\User_Manual_Remaining_Content.docx"
md_path = r"C:\Users\Lenovo\.gemini\antigravity\brain\6f428131-8d44-4151-b39a-94e60cd42594\artifacts\user_manual_remaining.md"

doc = Document(template_path)

# Clear existing content to only give the user the new content 
# (keeps styles, headers/footers, and margins intact)
for p in doc.paragraphs:
    p._element.getparent().remove(p._element)

with open(md_path, 'r', encoding='utf-8') as f:
    lines = f.readlines()

def safe_add_paragraph(doc, text, style_name):
    # fallback style logic
    styles = [s.name for s in doc.styles]
    p = doc.add_paragraph()
    
    if style_name in styles:
        p.style = doc.styles[style_name]
    else:
        matched = False
        if "Heading 3" == style_name:
            for s in styles:
                if s.lower() == "heading 3": p.style = doc.styles[s]; matched = True; break
        elif "Heading 4" == style_name:
            for s in styles:
                if s.lower() == "heading 4": p.style = doc.styles[s]; matched = True; break
        elif "List" in style_name:
             for s in styles:
                if "list paragraph" == s.lower() or "list bullet" == s.lower(): 
                   p.style = doc.styles[s]; matched = True; break
    
    # Process inline bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        else:
            p.add_run(part)
    return p

for line in lines:
    line = line.strip()
    if not line:
        continue
    
    if line.startswith('### '):
        safe_add_paragraph(doc, line.replace('### ', ''), 'Heading 3')
    elif line.startswith('#### '):
        safe_add_paragraph(doc, line.replace('#### ', ''), 'Heading 4')
    elif line.startswith('**Step'):
        # Just bold the text
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
    elif line.startswith('*Figure'):
        p = doc.add_paragraph()
        run = p.add_run(line.replace('*', ''))
        run.italic = True
        p.alignment = 1 # Center
    elif line.startswith('* '):
        safe_add_paragraph(doc, line[2:], 'List Paragraph')
    else:
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)

doc.save(out_path)
print("Successfully saved to", out_path)
