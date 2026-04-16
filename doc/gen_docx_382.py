"""
gen_docx_382.py
Generates doc/Report 6/3.2.8_AI_based_Flood_Prediction.docx
formatted exactly to match the Report 6 User Manual template.
Run: python gen_docx_382.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

# ── paths ──────────────────────────────────────────────────────────────────────
TEMPLATE_PATH = r"d:\FPTU\SEM9\SEP490\FE\FDA-Web\doc\Template User Manual Report 6.docx"
OUT_DIR       = r"d:\FPTU\SEM9\SEP490\FE\FDA-Web\doc\Report 6"
OUT_PATH      = os.path.join(OUT_DIR, "3.2.8_AI_based_Flood_Prediction.docx")

os.makedirs(OUT_DIR, exist_ok=True)

# ── helpers ────────────────────────────────────────────────────────────────────

def copy_style_from_para(src_para, dst_para):
    """Copy font/styling from src paragraph to dst paragraph runs."""
    for dst_run in dst_para.runs:
        src_fmt = src_para.paragraph_format
        try:
            dst_run.font.name = src_run.font.name if src_run.font.name else 'Calibri'
        except Exception:
            pass


def find_style(doc: Document, name_like: str):
    """Return the first style name in doc whose lowercase contains name_like."""
    for s in doc.styles:
        if name_like.lower() in s.name.lower():
            return s.name
    return None


def add_heading(doc, text, style_name):
    p = doc.add_paragraph()
    p.style = doc.styles[style_name]
    p.add_run(text)
    return p


def add_normal(doc, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_step(doc, number_str, text):
    """Format a Step heading (bold number + text)."""
    p = doc.add_paragraph()
    run_num = p.add_run(number_str)
    run_num.bold = True
    run_text = p.add_run(text)
    run_text.bold = False
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.add_run(text)
    return p


def add_figure_caption(doc, number, description):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure {number}: {description}")
    run.italic = True
    run.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, col_widths=None):
    """
    Build a Word table.
    headers  : list of str
    rows      : list of list of str
    col_widths: list of float (cm), one per column; None = equal width
    """
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'

    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = str(val)

    # column widths
    if col_widths:
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci])

    return table


# ── document ───────────────────────────────────────────────────────────────────

doc = Document(TEMPLATE_PATH)

# ── 3.2.8 AI-based Flood Prediction ───────────────────────────────────────────
add_heading(doc, "3.2.8 AI-based Flood Prediction", "Heading 1")

p_intro = doc.add_paragraph(
    "This section describes how to use the AI-powered flood prediction feature "
    "in the FDA System. The feature consists of three main capabilities:"
)
add_bullet(doc, "View AI-based Flood Prediction – Run an AI analysis for a specific administrative area.")
add_bullet(doc, "Compare AI Predictions with Actual Data – Evaluate prediction accuracy by overlaying AI forecast windows against real-time water level data.")
add_bullet(doc, "Identify High-risk Flood Zones (AI) – Visually locate areas with the highest flood risk on the map.")

doc.add_paragraph()  # spacer

# ── 3.2.8.1 ───────────────────────────────────────────────────────────────────
add_heading(doc, "3.2.8.1 View AI-based Flood Prediction", "Heading 2")

doc.add_paragraph(
    "This feature allows authorized users (Admin, Authority, Super Admin) to trigger an "
    "AI analysis for any administrative area displayed on the flood map. The system computes "
    "a multi-source ensemble prediction and returns flood probability, confidence, estimated "
    "impact, and contributing sensor stations."
)

# ── Step 1 ──
add_step(doc, "Step 1: ", "Navigate to the Flood Map")
add_normal(doc,
    "Log in to the FDA System and access the Admin Dashboard. "
    "In the left-hand sidebar, click \"Zones\" (or \"Flood Map\", depending on your deployment). "
    "The map view loads, displaying administrative area polygons over a base map layer.")

# ── Step 2 ──
add_step(doc, "Step 2: ", "Confirm Required Map Layers are Active")
add_normal(doc,
    "Before running AI prediction, ensure the following layers are enabled via the "
    "Layer Panel:")
add_bullet(doc, "Khu Vực Admin – Toggle ON to show administrative area boundaries. Clicking an area triggers the prediction card.")
add_bullet(doc, "Trạm Quan Trắc – Toggle ON to display sensor station markers. Station data contributes to the AI prediction.")
add_normal(doc,
    'Open the Layer Panel by clicking the Layers button on the map interface. '
    'Adjust toggles as needed, then click "Lưu" (Save) to apply changes.')

# ── Step 3 ──
add_step(doc, "Step 3: ", "Select an Administrative Area on the Map")
add_normal(doc,
    'Click directly on any administrative area polygon (ward or district boundary) displayed on the map. '
    'The system highlights the selected area by zooming the map camera to fit the area\'s boundary. '
    'The Area Detail Card slides in from the left side of the screen, showing:')
add_bullet(doc, "Area name and administrative code")
add_bullet(doc, 'Two action buttons: "Chụp ảnh vệ tinh" (Satellite) and "Dự đoán AI" (AI Prediction)')

# ── Step 4 ──
add_step(doc, "Step 4: ", "Trigger AI Flood Prediction")
add_normal(doc,
    'Click the "Dự đoán AI" button (purple/violet button with a sparkles icon). '
    'The button label changes to "Đang gọi AI…" and becomes disabled while the request is in progress. '
    'The system calls the backend API to assemble data from multiple sources (weather data, satellite imagery, '
    'sensor stations, and terrain analysis) and compute the ensemble prediction.')

# ── Step 5 ──
add_step(doc, "Step 5: ", "Review the AI Prediction Result")
add_normal(doc,
    "Upon successful completion, the PredictFloodAiPanel section expands within the Area Detail Card, displaying the following subsections:")

add_heading(doc, "5.1 Model Result Hero Card", "Heading 3")
add_normal(doc, "The top section of the prediction panel displays a color-coded hero card:")
add_table(doc,
    ["Field", "Vietnamese Label", "Description"],
    [
        ["Flood Probability", "Xác suất lũ", "The ensemble model's predicted probability of flood occurrence, expressed as a percentage."],
        ["Risk Level", "Mức rủi ro", "Derived from probability: Cao (>70%), Trung bình (50–70%), Vừa (15–50%), Thấp (≤15%)."],
        ["Confidence", "Độ tin cậy", "The model's confidence score for this prediction, expressed as a percentage."],
    ],
    col_widths=[3.0, 3.5, 8.5]
)
doc.add_paragraph()
add_normal(doc,
    "The background color of the hero card reflects the risk level:")
add_bullet(doc, "Green gradient – Safe / Low risk")
add_bullet(doc, "Yellow gradient – Caution / Medium risk")
add_bullet(doc, "Orange gradient – Warning / Moderate risk")
add_bullet(doc, "Red gradient – Critical / High risk")

add_heading(doc, "5.2 Estimated Accuracy by Model Phase", "Heading 3")
add_normal(doc,
    'A section titled "Độ chính xác ước tính (theo giai đoạn mô hình)" shows the accuracy of each ensemble component:')
add_table(doc,
    ["Phase Key", "Source", "Description"],
    [
        ["baseline_accuracy", "WeatherAPI", "Baseline weather forecast accuracy"],
        ["phase_1_accuracy", "Open-Meteo", "Regional weather accuracy by coordinates"],
        ["phase_2_accuracy", "Copernicus", "Soil saturation / water absorption data"],
        ["phase_3_accuracy", "DEM", "Terrain slope, drainage, flood risk mapping"],
        ["total_improvement", "Full Stack", "Combined accuracy across all signal sources"],
    ],
    col_widths=[3.5, 2.5, 9.0]
)

add_heading(doc, "5.3 Impact Estimation", "Heading 3")
add_normal(doc, "A white card below the hero shows estimated flood impact:")
add_table(doc,
    ["Field", "Vietnamese Label", "Description"],
    [
        ["Estimated Depth", "Độ sâu nước ước tính", "Estimated flood water depth in meters"],
        ["Estimated Affected Area", "Phạm vi ảnh hưởng ước tính", "Estimated geographic area affected by the predicted flood"],
        ["Recommendation", "Khuyến nghị", "AI-generated recommended action in a color-matched recommendation box"],
    ],
    col_widths=[3.0, 4.0, 8.0]
)

add_heading(doc, "5.4 Area Status Summary", "Heading 3")
add_normal(doc,
    'A compact bar below the hero card shows "Trạng thái khu vực:" followed by the current status label '
    "(e.g., NORMAL, CAUTION, WARNING, CRITICAL) and an optional severity level badge (mức 1–4).")

add_heading(doc, "5.5 Administrative Area Information", "Heading 3")
add_normal(doc,
    'A card titled "Khu vực đánh giá" displays the area name, administrative code, and parent area of the selected district/ward.')

add_heading(doc, "5.6 Forecast Windows – Time-based Predictions", "Heading 3")
add_normal(doc,
    'A section titled "Xu hướng theo khung giờ" displays prediction chips for multiple forecast horizons '
    "(e.g., +1h, +3h, +5h, +7h). Each chip shows the forecast time label, flood probability percentage, "
    "and status (NORMAL / WARNING / CRITICAL). Chips are color-coded green/yellow/orange/red. "
    "This provides a forward-looking view of how flood risk evolves over the next several hours.")

add_heading(doc, "5.7 Monitoring Stations Near the Area", "Heading 3")
add_normal(doc,
    'A collapsible section titled "Trạm đo gần khu vực" (expandable by clicking) lists all sensor stations '
    "contributing data to the prediction. Each entry shows: station code and name, current water level in meters, "
    "severity badge (Nghiêm trọng / Cao / Trung bình / Thấp / Bình thường), and street/ward location.")

add_heading(doc, "5.8 Community Reports", "Heading 3")
add_normal(doc,
    'A collapsible section titled "Phản ánh từ người dân" (expandable by clicking) lists citizen-submitted '
    "flood reports within the area, each showing the report description, severity level badge, and timestamp.")

add_heading(doc, "5.9 AI Analysis & Recommendation", "Heading 3")
add_normal(doc,
    'A section titled "Phân tích & khuyến nghị (AI)" displays the AI consultant\'s full summary text '
    "in a color-matched recommendation box.")

add_heading(doc, "5.10 Technical Details (Expandable)", "Heading 3")
add_normal(doc,
    'A collapsible "Chi tiết kỹ thuật (thời tiết, đất, địa hình…)" section shows raw technical components '
    "(weather, saturation, terrain, historical similarity) as expandable JSON previews for advanced users.")

# ── Step 6 ──
add_step(doc, "Step 6: ", "Handle Errors and Mismatches")
add_normal(doc,
    "If the API call fails, an error banner appears below the action buttons inside the Area Detail Card:")
add_bullet(doc, "Vệ tinh error – Red banner with alert icon showing the error message")
add_bullet(doc, "Dự đoán AI error – Amber/yellow banner with alert icon showing the error message")
add_normal(doc,
    'If the returned area does not match the selected area on the map, a warning line appears: '
    '"Phản hồi dự đoán AI không khớp khu vực đang chọn trên bản đồ." '
    "(AI prediction response does not match the currently selected area on the map.)")

# ── Step 7 ──
add_step(doc, "Step 7: ", "Close the Area Detail Card")
add_normal(doc,
    'Click the × button at the top-right corner of the Area Detail Card to close it and deselect the area. '
    "Alternatively, click anywhere on the map background to deselect.")

doc.add_paragraph()

# ── 3.2.8.2 ───────────────────────────────────────────────────────────────────
add_heading(doc, "3.2.8.2 Compare AI Predictions with Actual Data", "Heading 2")

doc.add_paragraph(
    "This feature enables authorized users to evaluate the accuracy of AI flood predictions by "
    "comparing the predicted forecast windows against real-time water level data from sensor stations. "
    "The comparison is facilitated through the Flood History module."
)

add_step(doc, "Step 1: ", "Navigate to the Flood History Module")
add_normal(doc,
    'From the Admin Dashboard sidebar, click "Flood History" or "Lịch Sử Ngập". '
    "The Flood History page loads with a Filter Bar (sticky on the left), KPI Cards row, and chart area.")

add_step(doc, "Step 2: ", "Open the Filter Panel")
add_normal(doc,
    "The Flood Filter Bar is expanded by default. If collapsed, click the filter panel header to expand it.")

add_step(doc, "Step 3: ", "Select Station(s) for Comparison")
add_normal(doc, "The filter bar supports two modes:")
add_bullet(doc, "Single Station Mode (default) – Use the Station dropdown to select one monitoring station. Select Trend Analysis as the View Mode.")
add_bullet(doc, 'Compare Mode (multi-station) – Toggle the "Compare Mode" switch to enable multi-station selection. Select up to 3 stations by clicking each station button. Selected stations display a checkmark. Comparison mode switches the view to Detailed History automatically.')

add_step(doc, "Step 4: ", "Choose a Time Period")
add_normal(doc, "In the Period dropdown, select one of the following presets:")
add_bullet(doc, "Last 24 Hours")
add_bullet(doc, "Last 7 Days")
add_bullet(doc, "Last 30 Days")
add_bullet(doc, "Last 90 Days")
add_bullet(doc, "Last 365 Days")
add_normal(doc,
    "Select a period that aligns with the AI prediction window you want to evaluate "
    "(e.g., select 'Last 7 Days' to compare against a 5-hour or 7-hour AI forecast).")

add_step(doc, "Step 5: ", "Apply the Filters")
add_normal(doc, 'Click the "Apply Filters" button. The system fetches historical water level data for the selected station(s) and time period, and the chart area updates accordingly.')

add_step(doc, "Step 6: ", "View the Comparison Chart")
add_normal(doc, "After applying filters, the main chart area displays the following views:")

add_heading(doc, "Trend Analysis View (single station)", "Heading 3")
add_normal(doc,
    'A line/area chart titled "Phân tích ngập" (Flood Analysis) shows:')
add_bullet(doc, "X-axis: Time periods (days or hours depending on the selected period and granularity)")
add_bullet(doc, "Y-axis: Water level in centimeters (cm)")
add_bullet(doc, 'Primary line: "Mực nước TB" (Average Water Level) – the actual observed data')
add_bullet(doc, 'Secondary line: "Mực nước max" (Maximum Water Level) – the peak observed level')
add_bullet(doc, "Flood hours indicator in chart tooltip")
add_normal(doc,
    'A comparison footer section shows comparison badges relative to the previous period: '
    '"Mực TB: ±X%" (average water level change) and "Giờ ngập: ±X%" (flood hours change). '
    "Green badge indicates improvement; red badge indicates worsening.")

add_heading(doc, "Detailed History View (compare mode or detailed view)", "Heading 3")
add_normal(doc,
    'A bar chart titled "Phân tích ngập" displays:')
add_bullet(doc, 'Grouped bars per time period (e.g., per day) showing "Mực max" (Maximum, red) and "Mực TB" (Average, primary color)')
add_bullet(doc, "Hover over a bar to see the exact value and flood hours for that period")
add_normal(doc, 'The chart tooltip indicates: "[value] cm · [X]h ngập"')

add_step(doc, "Step 7: ", "Interpret the Comparison with AI Forecast Windows")
add_normal(doc, "To evaluate AI prediction accuracy against actual data:")
add_bullet(doc, "Identify the forecast horizon of interest in the AI prediction (e.g., +3h, +5h, +7h forecast window from Step 5.6 in section 3.2.8.1).")
add_bullet(doc, "In the Flood History chart, locate the corresponding time point or period.")
add_bullet(doc, "Compare AI Forecast Probability (from PredictFloodAiPanel) vs Actual Flood Hours (from chart).")
add_bullet(doc, "Compare AI Risk Level (Cao / Trung bình / Vừa / Thấp) vs Actual Max Water Level (cm).")
add_bullet(doc, "A forecast window with high probability (>70%) that corresponds to a period with high actual flood hours indicates accurate prediction.")
add_bullet(doc, "Divergence between predicted risk and actual data may indicate model drift or need for re-calibration.")

add_step(doc, "Step 8: ", "Log Prediction Records")
add_normal(doc,
    "To maintain a training dataset for future model improvement, record the following for each comparison:")
add_table(doc,
    ["Field", "Description"],
    [
        ["Administrative area name", "Name of the area being evaluated"],
        ["Forecast timestamp", "Date and time the AI prediction was generated"],
        ["AI predicted probability and risk level", "Values from the PredictFloodAiPanel hero card"],
        ["Actual observed water level", "Max and average water level from the Flood History chart (cm)"],
        ["Actual flood hours", "Total flood hours for the corresponding period from the chart"],
        ["Deviation", "Difference between predicted and actual values"],
    ],
    col_widths=[5.0, 10.0]
)
doc.add_paragraph()
add_normal(doc,
    "This data supports model retraining and accuracy improvement over time when the system is "
    "deployed at scale.")

add_step(doc, "Step 9: ", "View KPI Summary")
add_normal(doc, "Above the chart, KPI Cards provide a quick statistical summary for the selected period:")
add_table(doc,
    ["KPI", "Vietnamese Label", "Metric", "Description"],
    [
        ["Max Water Level", "Mực nước max", "[X] cm", "Highest water level recorded in the selected period"],
        ["Avg Water Level", "Mực nước TB", "[X] cm", "Average water level across the period"],
        ["Total Flood Hours", "Giờ ngập", "[X]h", "Total hours where water level exceeded the warning threshold"],
        ["Data Completeness", "Độ đầy đủ dữ liệu", "[X]%", "Percentage of expected data points that were recorded"],
    ],
    col_widths=[2.5, 3.5, 2.0, 7.0]
)

doc.add_paragraph()

# ── 3.2.8.3 ───────────────────────────────────────────────────────────────────
add_heading(doc, "3.2.8.3 Identify High-risk Flood Zones (AI)", "Heading 2")

doc.add_paragraph(
    "This feature highlights areas with the highest flood risk on the interactive map using AI-driven "
    "probability data, enabling authorities to prioritize response and mitigation efforts."
)

add_step(doc, "Step 1: ", "Navigate to the Zones / Flood Map")
add_normal(doc,
    'From the Admin Dashboard, click "Zones" in the sidebar. '
    "The interactive map view loads with administrative area boundaries and sensor station markers.")

add_step(doc, "Step 2: ", "Enable the Flood Severity Layer")
add_normal(doc, 'Open the Layer Panel. Toggle "Trạm Quan Trắc" to ON to activate the flood severity overlay. '
    'Click "Lưu" (Save) to apply the change.')
add_normal(doc, "The map now displays:")
add_bullet(doc, "Sensor station markers colored by current severity level")
add_bullet(doc, "Critical radius rings – Large pulsing red circles around stations with Critical severity")

add_step(doc, "Step 3: ", "Reference the Flood Legend")
add_normal(doc,
    'The "Mức độ ngập" (Flood Severity) legend appears in the bottom-left corner of the map when the station layer is active:")
add_table(doc,
    ["Level", "Vietnamese Label", "Color", "Water Level Threshold"],
    [
        ["Safe", "An toàn", "Green (#16A34A)", "< 10 cm"],
        ["Caution", "Cảnh giác", "Yellow (#CA8A04)", "10–20 cm"],
        ["Warning", "Cảnh báo", "Orange (#EA580C)", "20–40 cm"],
        ["Critical", "Nghiêm trọng", "Red (#DC2626)", "≥ 40 cm (pulsing)"],
    ],
    col_widths=[2.0, 3.5, 2.5, 7.0]
)
doc.add_paragraph()

add_step(doc, "Step 4: ", "Identify High-risk Zones via AI Prediction")
add_normal(doc, "For each administrative area on the map:")
add_bullet(doc, "Click the area polygon to open the Area Detail Card.")
add_bullet(doc, 'Click "Dự đoán AI" to run the AI prediction (as described in section 3.2.8.1).')
add_bullet(doc, "Read the risk level from the hero card:")
add_bullet(doc, "Cao (red gradient) – Highest risk. Immediate attention required.", level=1)
add_bullet(doc, "Trung bình (orange gradient) – Moderate risk.", level=1)
add_bullet(doc, "Vừa (yellow gradient) – Medium risk.", level=1)
add_bullet(doc, "Thấp (green gradient) – Low risk.", level=1)
add_bullet(doc, 'Check the "Xu hướng theo khung giờ" (Forecast Windows) section to see how risk evolves over the next 3–7 hours. Windows marked CRITICAL indicate time-specific high-risk periods.')
add_bullet(doc, 'Check the "Contributing Stations" section to identify which specific stations are driving the high-risk assessment.')

add_step(doc, "Step 5: ", "Locate Critical Zones on the Map")
add_normal(doc, "High-risk zones are visually identifiable on the map through the following indicators:")
add_bullet(doc, "Red pulsing markers (critical stations) – Large animated red dots with pulsing outer ring at critical station locations.")
add_bullet(doc, "Red gradient hero cards in the Area Detail Card – Critical risk areas display a red gradient background for immediate visual recognition.")
add_bullet(doc, "Red window chips in the forecast timeline – Specific hours marked CRITICAL are shown in red chips.")
add_normal(doc, "To quickly find the highest-risk zones across the map:")
add_bullet(doc, "Toggle Trạm Quan Trắc layer ON.")
add_bullet(doc, "Visually scan the map for red pulsing markers.")
add_bullet(doc, 'Click each red marker to open the Flood Detail Card, showing: station name, current water level (cm), alert level (CRITICAL / WARNING / CAUTION / SAFE), sensor height, distance, and last update timestamp.')
add_bullet(doc, 'Click "Xem trạm" to navigate to the full station detail page.')

add_step(doc, "Step 6: ", "Respond to Identified High-risk Zones")
add_normal(doc, "Based on the AI-identified high-risk zones, authorities can take the following actions:")
add_table(doc,
    ["Risk Level", "Vietnamese Label", "Recommended Action"],
    [
        ["Critical", "Cao", "Immediate deployment of response teams. Issue alerts to affected residents. Monitor real-time water levels continuously."],
        ["Warning", "Trung bình", "Pre-position resources. Issue advisory notices. Continue enhanced monitoring."],
        ["Caution", "Vừa", "Enhanced monitoring. Prepare contingency plans."],
        ["Safe", "Thấp", "Normal operations. Continue routine monitoring."],
    ],
    col_widths=[2.0, 3.0, 10.0]
)
doc.add_paragraph()

add_step(doc, "Step 7: ", "Overlay Satellite Analysis (Optional)")
add_normal(doc,
    "For enhanced situational awareness, authorities can overlay satellite flood analysis on top of the station layer:")
add_bullet(doc, "Click an administrative area polygon to open the Area Detail Card.")
add_bullet(doc, 'Click "Chụp ảnh vệ tinh" (Satellite Analysis button).')
add_bullet(doc, 'The label changes to "Đang phân tích…" while the API call is in progress.')
add_bullet(doc, "Upon success, satellite-derived flood anomaly polygons are overlaid on the map within the selected area.")
add_bullet(doc, "This provides a visual confirmation of actual flood extent, complementing the AI prediction and station data.")

# ── save ───────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
